from docx import Document
import fitz
import os
import re
import unicodedata

def is_variant(character):
    category = unicodedata.category(character)
    # if category in ['Lu', 'Ll', 'Lo', 'Nd', 'Po', 'Ps', 'Pe', 'Pd', 'Zs', 'Sm']:
    if category not in ['So']:
        return False
    return True

def get_variant_character(text):
    t = []
    for char in text:
        if is_variant(char) == True:
            t.append(char)
    return t

def get_pdf_text(file_path):
    doc = fitz.Document(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return get_variant_character(text)

def get_word_text(file_path):
    # 打开Word文档
    doc = Document(file_path)

    text = ''
    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # 遍历文档中的每个表格
    for table in doc.tables:
        # 遍历表格中的每个单元格，并获取其文字内容
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    return get_variant_character(text)

def find_files(folder):
    arr = []
    for root_main, dirs_main, files_main in os.walk(folder):
        # 遍历当前文件夹内的子文件夹
        for sub_folder in dirs_main:
            for root, dirs, files in os.walk(os.path.join(root_main, sub_folder)):
                pdf_path = ''
                word_path = ''
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_path = os.path.join(root, file)
                    elif file.lower().endswith('.docx'):
                        word_path = os.path.join(root, file)
                if pdf_path:
                    print(pdf_path)
                    arr.extend(get_pdf_text(pdf_path))
                elif word_path:
                    print(word_path)
                    arr.extend(get_word_text(word_path))

    with open(os.path.join(folder, '1.txt'), 'w', encoding='utf-8') as f:
        f.write(''.join(arr))
    new_path = os.path.join(folder, '2.txt') 
    new_arr = to_python(arr, os.path.join(folder, '3.txt') )
    with open(new_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(new_arr))

def to_python(arr, new_path):
    dic = {}
    with open(new_path, 'r', encoding='utf-8') as f:
        txt = f.read()
    for t in txt.split('\n'):
        r = re.search(r'"(.*?)".*?"(.*?)"', t)
        if r:
            if r.group(1) and r.group(2):
                dic[r.group(1)] = ''

    new_arr = []
    for char in arr:
        if char in ['●', '○', '〓', '▲']:
            continue
        if char not in dic:
            dic[char] = 1
            new_arr.append(f's = s.replace("{char}", "")')
    
    return new_arr

def main():
    folder_path = r'D:\python\env\uiautobot\ExtractPDfInformation\data\已完成'
    find_files(folder_path)

if __name__ == '__main__':
    main()
